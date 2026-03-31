import streamlit as st
import json
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Cm
import io

# =============================
# CONFIG
# =============================
st.set_page_config(page_title="LADOS Gerador", layout="wide")

genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

# =============================
# LOAD JSON
# =============================
def carregar_json(caminho):
    with open(caminho, "r", encoding="utf-8") as f:
        return json.load(f)

# =============================
# PROMPT MESTRE
# =============================
def montar_prompt(ano, disciplina, descritor, habilidade, exemplos):
    exemplos_texto = "\n".join(exemplos)

    prompt = f"""
Você é um especialista em avaliação educacional no padrão SAEB.

OBJETIVO:
Criar 1 questão inédita alinhada ao descritor abaixo.

ANO: {ano}
DISCIPLINA: {disciplina}
DESCRITOR: {descritor}
HABILIDADE: {habilidade}

EXEMPLOS DE REFERÊNCIA:
{exemplos_texto}

REGRAS OBRIGATÓRIAS:
- 1 texto-base (se necessário)
- 4 alternativas (A, B, C, D)
- Apenas 1 correta
- Linguagem adequada ao ano

FORMATO:
Texto-base:
Pergunta:
A)
B)
C)
D)

Ao final, informe:
Gabarito:
Justificativa:
Descritor:
"""
    return prompt

# =============================
# VALIDAÇÃO
# =============================
def validar_questao(texto):
    criterios = ["A)", "B)", "C)", "D)", "Gabarito"]
    return all(c in texto for c in criterios)

# =============================
# GERAÇÃO EM LOTE
# =============================
def gerar_questoes_lote(model, prompt, n=5):
    questoes = []
    for _ in range(n):
        resposta = model.generate_content(prompt).text
        if validar_questao(resposta):
            questoes.append(resposta)
    return questoes

# =============================
# DOCX
# =============================
def gerar_docx(lista_questoes):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1.3)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    doc.add_heading('Simulado LADOS', 0)

    for i, q in enumerate(lista_questoes, 1):
        doc.add_paragraph(f"Questão {i}")
        doc.add_paragraph(q)
        doc.add_paragraph("\n")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# =============================
# UI
# =============================
st.title("🎯 Gerador de Questões SAEB - LADOS")

arquivos = {
    "5º Ano - Português": "5ano_portugues.json",
    "5º Ano - Matemática": "5ano_matematica.json",
    "9º Ano - Português": "9ano_portugues.json",
    "9º Ano - Matemática": "9ano_matematica.json",
}

escolha = st.selectbox("Selecione Ano e Disciplina", list(arquivos.keys()))

dados = carregar_json(arquivos[escolha])

# Suporte a JSON simples ou otimizado
if "descritores" in dados:
    descritores = list(dados["descritores"].keys())
    descritor = st.selectbox("Descritor", descritores)
    habilidade = dados["descritores"][descritor]
    exemplos = []
else:
    descritores = list(dados.keys())
    descritor = st.selectbox("Descritor", descritores)
    habilidade = dados[descritor]["habilidade"]
    exemplos = dados[descritor].get("exemplos", [])

quantidade = st.slider("Quantidade de questões", 1, 10, 5)

if st.button("🚀 Gerar Questões"):
    with st.spinner("O Gemini está analisando os padrões e gerando os itens..."):
        try:
            # Limpeza dos exemplos
            if isinstance(exemplos, list):
                lista_limpa = [str(ex) for ex in exemplos if ex]
            else:
                lista_limpa = []

            # Criar modelo
            model = genai.GenerativeModel("gemini-1.5-flash-latest")

            # Prompt correto
            prompt = montar_prompt(
                ano=escolha,
                disciplina="",
                descritor=descritor,
                habilidade=habilidade,
                exemplos=lista_limpa
            )

            # Gerar questões
            questoes = gerar_questoes_lote(model, prompt, quantidade)

            if questoes:
                st.success(f"{len(questoes)} questões geradas com sucesso!")

                for i, q in enumerate(questoes, 1):
                    with st.expander(f"Questão {i}", expanded=True):
                        st.write(q)

                # DOCX correto
                docx_bytes = gerar_docx(questoes)

                st.download_button(
                    label="📥 Baixar Simulado em Word",
                    data=docx_bytes,
                    file_name=f"LADOS_{descritor}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("Nenhuma questão válida foi gerada. Tente novamente.")

        except Exception as e:
            st.error(f"Ocorreu um erro na geração: {e}")
